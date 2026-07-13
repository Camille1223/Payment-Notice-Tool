import os
import io
import re
import zipfile
import tempfile
import threading
import uuid
import time
from copy import deepcopy

from flask import Flask, request, send_file, render_template_string, jsonify, Response, stream_with_context
from docx import Document
from docx.oxml.ns import qn
import openpyxl

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB

# ──────────────────────────────────────────────
# In-memory job store for progress tracking
# ──────────────────────────────────────────────
_jobs: dict[str, dict] = {}
_jobs_lock = threading.Lock()


def job_update(job_id: str, **kw):
    with _jobs_lock:
        if job_id in _jobs:
            _jobs[job_id].update(kw)


HTML = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>付款通知书生成器</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    background: #f0f2f5;
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    padding: 24px;
  }
  .card {
    background: #fff;
    border-radius: 12px;
    box-shadow: 0 2px 16px rgba(0,0,0,0.1);
    padding: 40px;
    width: 100%;
    max-width: 560px;
  }
  .logo { font-size: 26px; font-weight: 700; color: #1a73e8; margin-bottom: 6px; }
  .subtitle { color: #666; font-size: 14px; margin-bottom: 28px; }
  .upload-area {
    border: 2px dashed #d0d7de;
    border-radius: 8px;
    padding: 20px;
    text-align: center;
    cursor: pointer;
    transition: border-color 0.2s, background 0.2s;
    margin-bottom: 14px;
    position: relative;
  }
  .upload-area:hover, .upload-area.drag-over {
    border-color: #1a73e8;
    background: #f6f9ff;
  }
  .upload-area input[type=file] {
    position: absolute; inset: 0; opacity: 0; cursor: pointer; width: 100%; height: 100%;
  }
  .upload-icon { font-size: 28px; margin-bottom: 6px; }
  .upload-label { font-size: 14px; color: #444; font-weight: 500; }
  .upload-hint { font-size: 12px; color: #999; margin-top: 3px; }
  .file-badge {
    display: none;
    margin-top: 8px;
    padding: 3px 10px;
    background: #e8f0fe;
    border-radius: 4px;
    font-size: 12px;
    color: #1a73e8;
    word-break: break-all;
  }
  .tip { background: #fffbe6; border: 1px solid #ffe58f; border-radius: 6px; padding: 10px 14px; font-size: 12px; color: #7d6608; margin-bottom: 14px; line-height: 1.6; }
  .tip code { background: #fff3d0; padding: 1px 5px; border-radius: 3px; font-family: monospace; }
  .btn {
    width: 100%;
    padding: 13px;
    background: #1a73e8;
    color: #fff;
    border: none;
    border-radius: 8px;
    font-size: 15px;
    font-weight: 600;
    cursor: pointer;
    margin-top: 6px;
    transition: background 0.2s;
  }
  .btn:hover:not(:disabled) { background: #1557b0; }
  .btn:disabled { background: #a8c7fa; cursor: not-allowed; }
  .progress-wrap { margin-top: 14px; display: none; }
  .progress-bar {
    height: 6px; background: #e8eaed; border-radius: 3px;
    overflow: hidden;
  }
  .progress-fill {
    height: 100%; background: #1a73e8; border-radius: 3px;
    transition: width 0.3s ease; width: 0%;
  }
  .progress-text { font-size: 12px; color: #666; margin-top: 5px; text-align: right; }
  .status {
    margin-top: 14px; padding: 11px 14px; border-radius: 8px;
    font-size: 13px; display: none; line-height: 1.5;
  }
  .status.info    { background: #e8f0fe; color: #1a73e8; }
  .status.success { background: #e6f4ea; color: #137333; }
  .status.error   { background: #fce8e6; color: #c5221f; }
  hr { border: none; border-top: 1px solid #eee; margin: 24px 0; }
</style>
</head>
<body>
<div class="card">
  <div class="logo">付款通知书生成器</div>
  <div class="subtitle">上传模板与数据表，自动生成 Word + PDF，打包下载</div>
  <hr>
  <div class="tip">
    <b>模板字段格式：</b>将 Word 模板中需要替换的内容改为 <code>{字段名}</code>，例如 <code>{公司名称}</code>、<code>{金额}</code>。
    Excel Sheet1 第一行列名需与模板字段名一致（不含花括号）。
  </div>

  <div class="upload-area" id="wordArea">
    <input type="file" id="wordFile" accept=".docx" onchange="setFile('word',this)">
    <div class="upload-icon">📄</div>
    <div class="upload-label">付款通知书模板 (.docx)</div>
    <div class="upload-hint">点击或拖拽上传</div>
    <div class="file-badge" id="wordName"></div>
  </div>

  <div class="upload-area" id="excelArea">
    <input type="file" id="excelFile" accept=".xlsx,.xls" onchange="setFile('excel',this)">
    <div class="upload-icon">📊</div>
    <div class="upload-label">付款通知数据表 (.xlsx)</div>
    <div class="upload-hint">使用 Sheet1，第一行为列名</div>
    <div class="file-badge" id="excelName"></div>
  </div>

  <button class="btn" id="btn" disabled onclick="generate()">生成并下载 ZIP</button>

  <div class="progress-wrap" id="progressWrap">
    <div class="progress-bar"><div class="progress-fill" id="fill"></div></div>
    <div class="progress-text" id="progressText">0%</div>
  </div>
  <div class="status" id="status"></div>
</div>
<script>
function setFile(type, input) {
  const f = input.files[0];
  if (!f) return;
  const badge = document.getElementById(type + 'Name');
  badge.textContent = f.name;
  badge.style.display = 'inline-block';
  document.getElementById('btn').disabled =
    !(document.getElementById('wordFile').files[0] && document.getElementById('excelFile').files[0]);
}

['wordArea','excelArea'].forEach(id => {
  const el = document.getElementById(id);
  el.addEventListener('dragover', e => { e.preventDefault(); el.classList.add('drag-over'); });
  ['dragleave','drop'].forEach(ev => el.addEventListener(ev, () => el.classList.remove('drag-over')));
});

function showStatus(msg, type) {
  const s = document.getElementById('status');
  s.innerHTML = msg;
  s.className = 'status ' + type;
  s.style.display = 'block';
}

function setProgress(pct, text) {
  document.getElementById('fill').style.width = pct + '%';
  document.getElementById('progressText').textContent = text || (pct + '%');
}

let _pollTimer = null;

function startPolling(jobId) {
  _pollTimer = setInterval(async () => {
    try {
      const r = await fetch('/progress/' + jobId);
      if (!r.ok) return;
      const d = await r.json();
      if (d.pct !== undefined) setProgress(d.pct, d.label);
      if (d.done) {
        clearInterval(_pollTimer);
        if (d.error) {
          showStatus('❌ ' + d.error, 'error');
          resetBtn();
        } else {
          setProgress(100, '完成');
          showStatus('✅ 生成成功！ZIP 已开始下载。', 'success');
          const a = document.createElement('a');
          a.href = '/download/' + jobId;
          a.download = '付款通知书.zip';
          a.click();
          resetBtn();
        }
      }
    } catch(e) {}
  }, 600);
}

function resetBtn() {
  const btn = document.getElementById('btn');
  btn.disabled = false;
  btn.textContent = '生成并下载 ZIP';
  setTimeout(() => { document.getElementById('progressWrap').style.display = 'none'; setProgress(0,''); }, 2000);
}

async function generate() {
  const btn = document.getElementById('btn');
  btn.disabled = true;
  btn.textContent = '生成中，请稍候…';
  document.getElementById('progressWrap').style.display = 'block';
  setProgress(5, '上传中…');
  showStatus('正在上传文件…', 'info');

  const form = new FormData();
  form.append('template', document.getElementById('wordFile').files[0]);
  form.append('data', document.getElementById('excelFile').files[0]);

  try {
    const resp = await fetch('/start', { method: 'POST', body: form });
    if (!resp.ok) {
      const err = await resp.json().catch(() => ({ error: resp.statusText }));
      throw new Error(err.error || '服务器错误');
    }
    const { job_id } = await resp.json();
    showStatus('正在生成文件，请稍候…', 'info');
    startPolling(job_id);
  } catch(e) {
    showStatus('❌ ' + e.message, 'error');
    resetBtn();
  }
}
</script>
</body>
</html>
"""


# ──────────────────────────────────────────────
# Excel parsing
# ──────────────────────────────────────────────

def excel_rows(xlsx_bytes: bytes) -> list[dict]:
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    ws = wb.active  # use active sheet instead of hardcoded "Sheet1"
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise ValueError("表格为空")
    headers = [str(h).strip() if h is not None else "" for h in rows[0]]
    data = []
    for row in rows[1:]:
        if not any(v is not None for v in row):
            continue
        record = {}
        for h, v in zip(headers, row):
            if h:
                record[h] = str(v).strip() if v is not None else ""
        data.append(record)
    return data


# ──────────────────────────────────────────────
# DOCX merge  –  handles {field} across split runs
# ──────────────────────────────────────────────

def _run_has_drawing(run) -> bool:
    """Return True if this run contains an image/drawing element (must not be touched)."""
    return bool(run._r.findall('.//' + qn('w:drawing')))


def _merge_runs_in_paragraph(para):
    """Coalesce consecutive text-only runs so split {field} placeholders can be matched.
    Runs that contain drawings are left completely alone."""
    runs = para.runs
    if len(runs) < 2:
        return

    # Build groups of consecutive text-only runs; flush each group into its first run.
    group_start = None
    group_text = ""
    for run in runs:
        if _run_has_drawing(run):
            # flush any open group, then skip this run entirely
            if group_start is not None:
                group_start.text = group_text
                group_start = None
                group_text = ""
            continue
        if group_start is None:
            group_start = run
            group_text = run.text or ""
        else:
            group_text += run.text or ""
            run.text = ""
    if group_start is not None:
        group_start.text = group_text


def _replace_in_paragraph(para, mapping: dict):
    _merge_runs_in_paragraph(para)
    for run in para.runs:
        if _run_has_drawing(run):
            continue
        for key, val in mapping.items():
            placeholder = "{" + key + "}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, val)


def _replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, mapping)
            for tbl in cell.tables:
                _replace_in_table(tbl, mapping)


def fill_template(template_bytes: bytes, mapping: dict) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)
    for table in doc.tables:
        _replace_in_table(table, mapping)
    # headers / footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for para in hdr.paragraphs:
                _replace_in_paragraph(para, mapping)
            for table in hdr.tables:
                _replace_in_table(table, mapping)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ──────────────────────────────────────────────
# PDF conversion  –  LibreOffice (preferred) or
#                    weasyprint (fallback)
# ──────────────────────────────────────────────

def _pdf_via_libreoffice(docx_path: str, out_dir: str) -> str | None:
    import subprocess
    try:
        r = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, timeout=30
        )
        if r.returncode == 0:
            base = os.path.splitext(os.path.basename(docx_path))[0]
            p = os.path.join(out_dir, base + ".pdf")
            return p if os.path.exists(p) else None
    except Exception:
        pass
    return None


def _pdf_via_weasyprint(docx_path: str, out_dir: str) -> str | None:
    """Convert docx → HTML → PDF using weasyprint (pure Python, CJK-capable)."""
    try:
        from docx2python import docx2python
        import html as html_lib
        from weasyprint import HTML as WH, CSS
        # Extract text as simple HTML via docx2python
        result = docx2python(docx_path)
        lines = []
        for body_section in result.body:
            for para in body_section:
                for cell in para:
                    text = " ".join(cell)
                    lines.append("<p>" + html_lib.escape(text) + "</p>")
        html_content = """<!DOCTYPE html><html><head>
<meta charset="UTF-8">
<style>
@font-face { font-family: 'NotoSansSC'; src: local('Noto Sans CJK SC'), local('Source Han Sans CN'), local('Microsoft YaHei'), local('SimHei'); }
body { font-family: 'NotoSansSC', 'Microsoft YaHei', SimHei, sans-serif; font-size: 11pt; margin: 2cm; line-height: 1.6; }
p { margin-bottom: 6pt; }
</style></head><body>""" + "\n".join(lines) + "</body></html>"
        pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        WH(string=html_content).write_pdf(pdf_path)
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        app.logger.warning("weasyprint fallback failed: %s", e)
        return None


def docx_to_pdf(docx_path: str, out_dir: str) -> str | None:
    # Try LibreOffice first (available in Docker image); fall back to weasyprint
    pdf = _pdf_via_libreoffice(docx_path, out_dir)
    if pdf:
        return pdf
    return _pdf_via_weasyprint(docx_path, out_dir)


# ──────────────────────────────────────────────
# Background worker
# ──────────────────────────────────────────────

def _run_job(job_id: str, template_bytes: bytes, data_bytes: bytes):
    try:
        rows = excel_rows(data_bytes)
    except Exception as e:
        job_update(job_id, done=True, error=f"读取 Excel 失败：{e}")
        return

    if not rows:
        job_update(job_id, done=True, error="Sheet1 没有数据行")
        return

    total = len(rows)
    name_count: dict[str, int] = {}
    zip_buffer = io.BytesIO()

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_count = 0
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, row in enumerate(rows, 1):
                company = row.get("公司名称", f"客户{i}")
                base = safe_filename(company) + "_付款通知书"
                name_count[base] = name_count.get(base, 0) + 1
                cnt = name_count[base]
                unique = base if cnt == 1 else f"{base}_{cnt}"

                # Generate DOCX
                try:
                    docx_bytes = fill_template(template_bytes, row)
                except Exception as e:
                    job_update(job_id, done=True, error=f"第 {i} 行填充模板失败：{e}")
                    return

                docx_path = os.path.join(tmpdir, f"{unique}.docx")
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                zf.write(docx_path, f"Word/{unique}.docx")

                # Generate PDF
                pdf_path = docx_to_pdf(docx_path, tmpdir)
                if pdf_path:
                    zf.write(pdf_path, f"PDF/{unique}.pdf")
                    pdf_count += 1

                pct = int(10 + 85 * i / total)
                job_update(job_id, pct=pct, label=f"已处理 {i}/{total}…")

    zip_buffer.seek(0)
    pdf_note = f"，其中 PDF {pdf_count} 份" if pdf_count else "（PDF 生成失败，仅含 Word）"
    job_update(job_id, pct=100, label="完成", done=True,
               zip_data=zip_buffer.getvalue(),
               summary=f"共生成 {total} 份{pdf_note}")


def safe_filename(name: str) -> str:
    return "".join(c for c in name if c not in r'\/:*?"<>|').strip() or "未知"


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/start", methods=["POST"])
def start():
    if "template" not in request.files or "data" not in request.files:
        return jsonify(error="缺少上传文件"), 400

    template_bytes = request.files["template"].read()
    data_bytes = request.files["data"].read()

    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {"pct": 5, "label": "已收到文件…", "done": False}

    t = threading.Thread(target=_run_job, args=(job_id, template_bytes, data_bytes), daemon=True)
    t.start()
    return jsonify(job_id=job_id)


@app.route("/progress/<job_id>")
def progress(job_id: str):
    with _jobs_lock:
        job = dict(_jobs.get(job_id, {}))
    if not job:
        return jsonify(error="job not found"), 404
    # don't send raw zip data over progress endpoint
    job.pop("zip_data", None)
    return jsonify(job)


@app.route("/download/<job_id>")
def download(job_id: str):
    with _jobs_lock:
        job = _jobs.get(job_id)
    if not job or not job.get("done"):
        return jsonify(error="未就绪"), 404
    if job.get("error"):
        return jsonify(error=job["error"]), 500
    zip_data = job.get("zip_data")
    if not zip_data:
        return jsonify(error="无数据"), 500
    return send_file(
        io.BytesIO(zip_data),
        mimetype="application/zip",
        as_attachment=True,
        download_name="付款通知书.zip",
    )


@app.route("/generate", methods=["POST"])
def generate_sync():
    """Legacy synchronous endpoint kept for backward compatibility."""
    if "template" not in request.files or "data" not in request.files:
        return jsonify(error="缺少上传文件"), 400
    template_bytes = request.files["template"].read()
    data_bytes = request.files["data"].read()

    try:
        rows = excel_rows(data_bytes)
    except Exception as e:
        return jsonify(error=f"读取 Excel 失败：{e}"), 400
    if not rows:
        return jsonify(error="Sheet1 没有数据行"), 400

    name_count: dict[str, int] = {}
    zip_buffer = io.BytesIO()
    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, row in enumerate(rows, 1):
                company = row.get("公司名称", f"客户{i}")
                base = safe_filename(company) + "_付款通知书"
                name_count[base] = name_count.get(base, 0) + 1
                cnt = name_count[base]
                unique = base if cnt == 1 else f"{base}_{cnt}"
                docx_bytes = fill_template(template_bytes, row)
                docx_path = os.path.join(tmpdir, f"{unique}.docx")
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                zf.write(docx_path, f"Word/{unique}.docx")
                pdf_path = docx_to_pdf(docx_path, tmpdir)
                if pdf_path:
                    zf.write(pdf_path, f"PDF/{unique}.pdf")

    zip_buffer.seek(0)
    return send_file(zip_buffer, mimetype="application/zip",
                     as_attachment=True, download_name="付款通知书.zip")


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    print("=" * 40)
    print("Dashboard: http://127.0.0.1:5000")
    print("=" * 40)
    app.run(debug=False, port=5000)
